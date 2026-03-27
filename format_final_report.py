from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\CyberGuard\ransomware-awareness-lab")
SRC_MD = ROOT / "COLLEGE_PROJECT_REPORT.md"
OUT_DOCX = ROOT / "CYBER_GUARD_FINAL_PROJECT_REPORT_WITH_CODE.docx"
ASSET_DIR = ROOT / "report-assets"

CODE_SNIPPETS = [
    (
        "Appendix A. Main Application Shell",
        ROOT / "src" / "App.tsx",
        120,
        "This code shows the overall application shell, navigation, authentication-aware entry flow, and theme toggle."
    ),
    (
        "Appendix B. Backend Server And Hash Lookup API",
        ROOT / "server.ts",
        180,
        "This code shows the Express server, MalwareBazaar hash lookup endpoint, and Socket.IO drill room handling."
    ),
    (
        "Appendix C. Malware Analyzer Logic",
        ROOT / "src" / "components" / "PhishingAnalyzer.tsx",
        170,
        "This code shows indicator types, suspicious component inspection, regex matching, and fallback analysis logic."
    ),
    (
        "Appendix D. Live Lab Realtime Drill",
        ROOT / "src" / "components" / "LiveLab.tsx",
        220,
        "This code shows the incident drill state, scenarios, telemetry behavior, and realtime attacker/defender simulation."
    ),
]

IMAGES = [
    ("01-login.png", "Figure 1. Login page with Google sign-in, demo mode, and theme toggle."),
    ("02-dashboard.png", "Figure 2. Dashboard overview showing the main modules of Cyber Guard."),
    ("03-about-ransomware.png", "Figure 3. About Ransomware module with business-of-extortion overview."),
    ("04-encryption-lab.png", "Figure 4. Encryption Lab showing the cryptographic learning interface."),
    ("05-malware-analyzer.png", "Figure 5. Malware Analyzer interface for suspicious indicator analysis."),
    ("06-threat-simulator.png", "Figure 6. Threat Simulator with attack profiles and ransom amount customization."),
    ("07-live-lab.png", "Figure 7. Live Lab showing staged compromise progression and training telemetry."),
    ("08-defense-center.png", "Figure 8. Defense Center presenting mitigation and response guidance."),
]


def set_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_cover_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("CYBER GUARD")
    set_font(r, size=24, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("RANSOMWARE AWARENESS LAB")
    set_font(r, size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("Final Project Report")
    set_font(r, size=16, italic=True)

    doc.add_paragraph("")
    doc.add_paragraph("")

    fields = [
        "Submitted By: ______________________________",
        "Register Number: ______________________________",
        "Department: ______________________________",
        "Institution: ______________________________",
        "Project Guide: ______________________________",
        "Academic Year: ______________________________",
    ]

    for field in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(field)
        set_font(r, size=13)

    doc.add_page_break()


def configure_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    for style_name, size in [("Normal", 12), ("Title", 20), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer_p)

    return doc


def render_markdown(doc, text):
    lines = text.splitlines()
    in_code = False
    code_lang = ""
    code_buffer = []

    def flush_code():
        nonlocal code_buffer, code_lang
        if not code_buffer:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        if code_lang:
            r = p.add_run(f"[{code_lang}]\n")
            set_font(r, name="Consolas", size=9, bold=True)
        r = p.add_run("\n".join(code_buffer))
        set_font(r, name="Consolas", size=9)
        code_buffer = []
        code_lang = ""

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
            else:
                flush_code()
                in_code = False
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.strip() == "---":
            p = doc.add_paragraph()
            p.add_run("_" * 68)
            continue

        match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if match:
            level = len(match.group(1))
            text_value = match.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(text_value, style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                p = doc.add_paragraph(text_value, style="Heading 1")
            elif level == 3:
                p = doc.add_paragraph(text_value, style="Heading 2")
            else:
                p = doc.add_paragraph(text_value, style="Heading 3")
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            p.paragraph_format.space_after = Pt(4)
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            continue

        if line.startswith("|") and line.endswith("|"):
            p = doc.add_paragraph()
            r = p.add_run(line)
            set_font(r, name="Consolas", size=9)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(line)
        set_font(r)

    flush_code()


def add_image_section(doc):
    doc.add_page_break()
    p = doc.add_paragraph("Visual Evidence", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for filename, caption in IMAGES:
        image_path = ASSET_DIR / filename
        if not image_path.exists():
            p = doc.add_paragraph()
            r = p.add_run(f"[Missing image: {filename}] {caption}")
            set_font(r, italic=True)
            continue

        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(6.2))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(caption)
        set_font(r, size=11, italic=True)


def add_code_appendix(doc):
    doc.add_page_break()
    p = doc.add_paragraph("Code Appendix", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for heading, path, max_lines, intro in CODE_SNIPPETS:
        p = doc.add_paragraph(heading, style="Heading 2")
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(4)
        r = meta.add_run(f"File: {path}")
        set_font(r, size=11, italic=True)

        intro_p = doc.add_paragraph()
        intro_p.paragraph_format.space_after = Pt(6)
        r = intro_p.add_run(intro)
        set_font(r, size=11)

        if not path.exists():
            missing = doc.add_paragraph()
            r = missing.add_run("[File not found]")
            set_font(r, italic=True)
            continue

        lines = path.read_text(encoding="utf-8").splitlines()[:max_lines]
        code_p = doc.add_paragraph()
        code_p.paragraph_format.space_after = Pt(8)
        r = code_p.add_run("\n".join(lines))
        set_font(r, name="Consolas", size=8)


if not SRC_MD.exists():
    raise FileNotFoundError(SRC_MD)

doc = configure_doc()
add_cover_page(doc)
render_markdown(doc, SRC_MD.read_text(encoding="utf-8"))
add_image_section(doc)
add_code_appendix(doc)
doc.save(OUT_DOCX)
print(OUT_DOCX)
