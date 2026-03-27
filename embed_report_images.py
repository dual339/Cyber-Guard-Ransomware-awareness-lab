from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(r"C:\CyberGuard\ransomware-awareness-lab")
REPORT_PATH = ROOT / "COLLEGE_PROJECT_REPORT.docx"
OUTPUT_PATH = ROOT / "COLLEGE_PROJECT_REPORT_WITH_IMAGES.docx"
ASSET_DIR = ROOT / "report-assets"

IMAGES = [
    ("01-login.png", "Figure 1. Login page with Google sign-in, demo mode, and theme toggle."),
    ("02-dashboard.png", "Figure 2. Dashboard overview showing the main modules of Cyber Guard."),
    ("03-about-ransomware.png", "Figure 3. About Ransomware module with business-of-extortion overview."),
    ("04-encryption-lab.png", "Figure 4. Encryption Lab showing the cryptographic learning interface."),
    ("05-malware-analyzer.png", "Figure 5. Malware Analyzer interface for suspicious indicator analysis."),
    ("06-threat-simulator.png", "Figure 6. Threat Simulator with attack profiles and ransom amount customization."),
    ("07-live-lab.png", "Figure 7. Live Lab showing staged compromise progression and training telemetry."),
    ("08-defense-center.png", "Figure 8. Defense Center presenting mitigation and response guidance."),
]

if not REPORT_PATH.exists():
    raise FileNotFoundError(f"Base report not found: {REPORT_PATH}")

doc = Document(REPORT_PATH)
styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
styles["Normal"].font.size = Pt(12)

doc.add_page_break()
heading = doc.add_paragraph()
heading.style = doc.styles["Heading 1"]
heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
heading.add_run("Visual Evidence")

for filename, caption in IMAGES:
    image_path = ASSET_DIR / filename
    if not image_path.exists():
      p = doc.add_paragraph()
      p.add_run(f"[Missing image: {filename}] ").bold = True
      p.add_run(caption)
      continue

    doc.add_paragraph("")
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(6.2))

    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    caption_run = caption_paragraph.add_run(caption)
    caption_run.italic = True
    caption_run.font.name = "Times New Roman"
    caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption_run.font.size = Pt(11)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
